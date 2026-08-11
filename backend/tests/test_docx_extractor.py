"""Verification: DOCX extractor.

Creates a test DOCX with headings, paragraphs, and a table
using python-docx, then extracts and verifies.

Run:
    python -m backend.tests.test_docx_extractor
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from docx import Document
from docx.shared import Inches

PASS = "[PASS]"
FAIL = "[FAIL]"
TEST_DIR = Path("data/processed/_test")
SOURCE_DOC_ID = "test-docx-00000000-0000-0000-0000-000000000000"


def _create_test_docx() -> Path:
    """Create a test DOCX with headings, paragraphs, a table, and an image."""
    TEST_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = TEST_DIR / "test_document.docx"

    doc = Document()
    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_paragraph(
        "This is the introductory paragraph for the test document. "
        "It contains information about machine learning and data science."
    )

    doc.add_heading("Section 1.1: Background", level=2)
    doc.add_paragraph(
        "Machine learning has become an essential tool in modern computing. "
        "This section discusses the historical background."
    )

    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Algorithm"
    table.cell(0, 1).text = "Use Case"
    table.cell(1, 0).text = "Random Forest"
    table.cell(1, 1).text = "Classification"
    table.cell(2, 0).text = "K-Means"
    table.cell(2, 1).text = "Clustering"

    doc.add_heading("Chapter 2: Methods", level=1)
    doc.add_paragraph(
        "Various machine learning methods are explored in this chapter."
    )

    # Add a small embedded image
    from PIL import Image
    import io

    img = Image.new("RGB", (80, 60), color=(100, 150, 200))
    img_bytes = io.BytesIO()
    img.save(img_bytes, format="PNG")
    img_bytes.seek(0)
    doc.add_picture(img_bytes, width=Inches(1.0))

    doc.save(str(docx_path))
    return docx_path


def test_docx_extraction():
    """Extract text blocks and images from the test DOCX."""
    from backend.processors.docx_extractor import extract_docx

    docx_path = _create_test_docx()
    result = extract_docx(docx_path, SOURCE_DOC_ID, "test_document.docx")

    has_headings = any(b.block_type == "heading" for b in result.text_blocks)
    has_paragraphs = any(b.block_type == "paragraph" for b in result.text_blocks)
    has_tables = any(b.block_type == "table" for b in result.text_blocks)
    has_images = len(result.extracted_images) >= 1
    has_full_text = len(result.full_text) > 100

    print(f"\n{'='*60}")
    print(f"Test: DOCX extraction")
    print(f"  Source document : {result.source_document_id}")
    print(f"  Text blocks     : {len(result.text_blocks)}")
    for b in result.text_blocks:
        label = f"[{b.block_type}]"
        if b.heading_level:
            label += f" (level {b.heading_level})"
        print(f"    {b.block_index}. {label} {b.text[:60]}...")
    print(f"  Has headings    : {has_headings}")
    print(f"  Has paragraphs  : {has_paragraphs}")
    print(f"  Has tables      : {has_tables}")
    print(f"  Extracted images: {len(result.extracted_images)}")
    for img in result.extracted_images:
        print(f"    - {img.image_path} (ext: {img.image_extension})")
    print(f"  Full text length: {len(result.full_text)} chars")
    print(f"  Warnings        : {len(result.warnings)}")

    ok = has_headings and has_paragraphs and has_tables and has_full_text
    print(f"  Result          : {PASS if ok else FAIL}")
    return ok


if __name__ == "__main__":
    print("=" * 60)
    print("DOCX EXTRACTOR VERIFICATION")
    print("=" * 60)

    results = [test_docx_extraction()]

    print(f"\n{'='*60}")
    passed = sum(results)
    print(f"Results: {passed}/{len(results)} tests passed")
    print("=" * 60)
    sys.exit(0 if all(results) else 1)
